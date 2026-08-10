from pathlib import Path
import re, shutil, zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

import build_week1_print_packet as pdf_builder

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "01_Curriculum" / "Culinary_1" / "Week_01_2026"
PKG = ROOT / "output" / "school_upload" / "Culinary_1_Week_01"
ZIP = ROOT / "output" / "school_upload" / "Culinary_1_Week_01_School_Upload.zip"
ORIGINAL_RECIPE = Path(r"C:\Users\blyth\Downloads\Pancakes (1).docx")

TEACHER = [
    "Day_01_Welcome_and_Kitchen_Orientation_DTG.md",
    "Day_02_Kitchen_Mode_and_Basic_Routines_DTG.md",
    "Day_03_Read_Before_You_Cook_DTG.md",
    "Day_04_First_Lab_Pancakes_DTG.md",
]
STUDENTS = [
    "Day_01_Kitchen_Orientation_Challenge.md",
    "Day_02_Kitchen_Mode_Quick_Practice.md",
    "Day_03_Pancake_Recipe_Reading_and_Preparation.md",
    "Day_04_Pancake_Lab_and_Reflection.md",
]

def clean(s):
    return (s.replace("â€œ", '"').replace("â€", '"').replace("â€™", "'")
             .replace("â€“", "-").replace("â€”", "-"))

def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pPr.append(shd)

def set_repeatable_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.05
    for name, size, color in [("Heading 1", 18, "17365D"), ("Heading 2", 13, "2F75B5"), ("Heading 3", 11, "17365D")]:
        st = doc.styles[name]; st.font.name = "Arial"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(8); st.paragraph_format.space_after = Pt(4); st.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        st = doc.styles[name]; st.font.name = "Arial"; st.font.size = Pt(10); st.paragraph_format.space_after = Pt(2)

def add_rich_text(p, text):
    parts = re.split(r"(\*\*.+?\*\*)", clean(text))
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else: p.add_run(part)

def recipe_content(doc):
    doc.add_heading("Instructor-Approved Pancake Recipe", level=2)
    doc.add_heading("Ingredients", level=3)
    for item in ["5 cups self-rising flour", "1/2 cup sugar", "2 cups milk", "2 cups buttermilk", "4 eggs", "4 Tbsp vegetable oil"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Method", level=3)
    for item in [
        "Preheat the griddle to medium.", "Whisk together the flour and sugar.",
        "In a separate bowl, whisk the milk, buttermilk, eggs, and oil.",
        "Add the wet ingredients to the dry and whisk for 10 seconds, or until the batter barely comes together. The batter will be lumpy.",
        "Use a #12 scoop (green) to portion batter onto the griddle.", "Cook until craters form.",
        "Do not add toppings for this Week 1 lab.", "When ready, flip the pancakes and cook until done."]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("Actual yield: __________ pancakes")

def add_markdown(doc, path, insert_recipe=False):
    inserted = False
    for raw in path.read_text(encoding="utf-8-sig").splitlines():
        line = raw.strip()
        if not line:
            continue
        if insert_recipe and not inserted and set(line) == {"_"} and len(line) > 20:
            recipe_content(doc); inserted = True; continue
        if insert_recipe and inserted and set(line) == {"_"} and len(line) > 20: continue
        if line.startswith("# "): doc.add_heading(clean(line[2:]), level=1)
        elif line.startswith("## "): doc.add_heading(clean(line[3:]), level=2)
        elif line.startswith("### "): doc.add_heading(clean(line[4:]), level=3)
        elif line.startswith("> "):
            if "Drafted for instructor review" not in line:
                p = doc.add_paragraph(); add_rich_text(p, line[2:]); p.style = doc.styles["Normal"]; shade_paragraph(p, "FFF2CC")
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number"); add_rich_text(p, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet"); add_rich_text(p, "____  " + line[6:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_rich_text(p, line[2:])
        elif set(line) == {"_"} and len(line) > 20:
            p = doc.add_paragraph("____________________________________________________________")
        else:
            p = doc.add_paragraph(); add_rich_text(p, line)

def add_cover(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70); p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Culinary 1"); r.font.name = "Arial"; r.font.size = Pt(14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Week 1 School Working Copy"); r.font.name = "Arial"; r.font.size = Pt(26); r.font.bold = False
    p = doc.add_paragraph("August 11-14, 2026 | Four instructional days"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading("Included", level=2)
    for x in ["Four Daily Teaching Guides", "Active projection source", "Instructor/TA observation and optional support tools", "One required print item: the group pancake recipe", "Instructor-approved pancake recipe source"]:
        doc.add_paragraph(x, style="List Bullet")
    p = doc.add_paragraph(); add_rich_text(p, "**School workflow:** Upload this DOCX to the work Google Drive and open it as a Google Doc when edits are needed. Print only the companion group-recipe PDF."); shade_paragraph(p, "DDEBF7")
    p = doc.add_paragraph(); add_rich_text(p, "**Recipe note:** The supplied pancake recipe is included. The Week 1 curriculum direction of no added toppings remains in force."); shade_paragraph(p, "FFF2CC")
    doc.add_page_break()

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.72); section.bottom_margin = Inches(0.62); section.left_margin = Inches(0.78); section.right_margin = Inches(0.78)
        hp = section.header.paragraphs[0]; hp.text = "Culinary 1 | Week 1 | August 11-14, 2026"; hp.style = doc.styles["Normal"]
        hp.runs[0].font.size = Pt(8); hp.runs[0].font.color.rgb = RGBColor(100,100,100)
        fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.add_run("Page "); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); fp._p.append(fld)

def build_docx(path):
    doc = Document(); set_repeatable_styles(doc); add_header_footer(doc); add_cover(doc)
    for f in TEACHER: add_markdown(doc, SRC/f); doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Instructional Supports"); r.font.name="Arial"; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string("17365D")
    doc.add_page_break()
    for f in STUDENTS:
        add_markdown(doc, SRC/"Student_Materials"/f)
        doc.add_page_break()
    add_markdown(doc, SRC/"Slides"/"Culinary_1_Week_01_Slide_Source.md")
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Required Print"); r.font.name="Arial"; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string("17365D")
    doc.add_page_break()
    add_markdown(doc, SRC/"Student_Materials"/"Day_03_Pancake_Recipe_Reading_and_Preparation.md")
    doc.core_properties.title = "Culinary 1 Week 1 School Working Copy"; doc.core_properties.author = "WACOS"
    doc.save(path)

def split_pdf(src, out, page_indexes):
    reader=PdfReader(src); writer=PdfWriter()
    for i in page_indexes: writer.add_page(reader.pages[i])
    with open(out,"wb") as f: writer.write(f)

def build_recipe_pdf(path):
    ss=getSampleStyleSheet()
    title=ParagraphStyle("title",parent=ss["Title"],fontName="Helvetica-Bold",fontSize=20,leading=23,textColor=colors.HexColor("#17365D"),spaceAfter=5)
    sub=ParagraphStyle("sub",parent=ss["BodyText"],fontName="Helvetica-Bold",fontSize=9,leading=11,spaceAfter=6)
    h=ParagraphStyle("h",parent=ss["Heading2"],fontName="Helvetica-Bold",fontSize=12,leading=14,textColor=colors.HexColor("#2F75B5"),spaceBefore=6,spaceAfter=3)
    body=ParagraphStyle("body",parent=ss["BodyText"],fontName="Helvetica",fontSize=9.5,leading=12,spaceAfter=2)
    bullet=ParagraphStyle("bullet",parent=body,leftIndent=15,firstLineIndent=-9)
    doc=SimpleDocTemplate(str(path),pagesize=letter,leftMargin=.72*inch,rightMargin=.72*inch,topMargin=.58*inch,bottomMargin=.55*inch,title="Culinary 1 Week 1 Pancake Recipe")
    story=[Paragraph("Pancake Recipe - Week 1 Group Copy",title),Paragraph("Culinary 1 | Thursday preparation and Friday lab.",sub),Paragraph("Actual yield: __________ pancakes",sub),Paragraph("Ingredients",h)]
    for x in ["5 cups self-rising flour","1/2 cup sugar","2 cups milk","2 cups buttermilk","4 eggs","4 Tbsp vegetable oil"]: story.append(Paragraph("-  "+x,bullet))
    story.append(Paragraph("Method",h))
    for i,x in enumerate(["Preheat the griddle to medium.","Whisk together the flour and sugar.","In a separate bowl, whisk the milk, buttermilk, eggs, and oil.","Add the wet ingredients to the dry and whisk for 10 seconds, or until the batter barely comes together. The batter will be lumpy.","Use a #12 scoop (green) to portion batter onto the griddle.","Cook until craters form.","Do not add toppings for this Week 1 lab.","When ready, flip the pancakes and cook until done."],1): story.append(Paragraph(f"{i}.  {x}",body))
    story += [Paragraph("Mark the Recipe",h),Paragraph("Underline ingredients. Box equipment. Circle measurements. Star action verbs. Draw an arrow to heat or doneness cues.",body),Spacer(1,4),Paragraph("Keep this copy with the group. Heat stays off until instructor approval.",sub)]
    doc.build(story)

def build():
    PKG.mkdir(parents=True, exist_ok=True)
    for f in PKG.iterdir():
        if f.is_file(): f.unlink()
    build_docx(PKG / "01_Culinary_1_Week_01_Editable.docx")
    build_recipe_pdf(PKG / "02_Culinary_1_Week_01_Required_Print_Pancake_Recipe.pdf")
    if ORIGINAL_RECIPE.exists(): shutil.copy2(ORIGINAL_RECIPE, PKG / "03_Pancakes_Recipe_Original.docx")
    (PKG/"README.txt").write_text(
        "CULINARY 1 - WEEK 1 SCHOOL UPLOAD PACKAGE\n\n"
        "01 Editable DOCX: upload/open as a Google Doc for day-to-day teacher planning.\n"
        "02 Required Print PDF: print one pancake recipe per group; reuse Thursday and Friday.\n"
        "03 Original Pancakes Recipe: supplied source recipe.\n\n"
        "The editable DOCX includes the active projection source, the instructor/TA orientation tool, optional re-entry practice, and the reusable Friday lab checklist.\n"
        "Only the pancake recipe is a required student print. Students record actual yield after the lab.\n"
        "The Week 1 no-toppings direction is retained.\n",
        encoding="utf-8")
    if ZIP.exists(): ZIP.unlink()
    with zipfile.ZipFile(ZIP,"w",zipfile.ZIP_DEFLATED) as z:
        for f in sorted(PKG.iterdir()): z.write(f, arcname=f"Culinary_1_Week_01/{f.name}")
    print(PKG); print(ZIP)

if __name__ == "__main__": build()
