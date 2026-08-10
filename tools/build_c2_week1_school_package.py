from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "01_Curriculum" / "Culinary_2" / "Week_01_2026"
PKG = ROOT / "output" / "school_upload" / "Culinary_2_Week_01"
ZIP = ROOT / "output" / "school_upload" / "Culinary_2_Week_01_School_Upload.zip"

TEACHER = [
    "Day_01_Reentry_Expectations_and_Recipe_Planning_DTG.md",
    "Day_02_Knife_Control_Baseline_DTG.md",
    "Day_03_Bistro_and_Sanitation_Readiness_DTG.md",
    "Day_04_Vegetable_Rice_Bowl_Baseline_Lab_DTG.md",
]

SUPPORTS = [
    "Day_01_Vegetable_Rice_Bowl_Recipe_and_Mise_Plan.md",
    "Day_02_Knife_Control_Baseline_Record.md",
    "Day_03_Bistro_and_Sanitation_Readiness_Check.md",
    "Day_04_Vegetable_Rice_Bowl_Lab_and_Reflection.md",
]


def shade_paragraph(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    properties.append(shade)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for name, size, color in [
        ("Heading 1", 18, "17365D"),
        ("Heading 2", 13, "2F75B5"),
        ("Heading 3", 11, "17365D"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)


def add_rich_text(paragraph, text):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def is_table_separator(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown(doc, path):
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                current = lines[index].strip()
                if not is_table_separator(current):
                    rows.append([cell.strip() for cell in current.strip("|").split("|")])
                index += 1
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            table.autofit = True
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    cell.text = value
                    for run in cell.paragraphs[0].runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8)
                        run.bold = row_index == 0
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("> "):
            paragraph = doc.add_paragraph()
            add_rich_text(paragraph, line[2:])
            shade_paragraph(paragraph, "FFF2CC")
        elif re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("- [ ] "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, "____  " + line[6:])
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, line[2:])
        else:
            paragraph = doc.add_paragraph()
            add_rich_text(paragraph, line)
        index += 1


def add_divider(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(120)
    run = paragraph.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string("17365D")
    doc.add_page_break()


def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        header = section.header.paragraphs[0]
        header.text = "Culinary 2 | Week 1 | August 11-14, 2026"
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Page ")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def build_docx(path):
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    run = title.add_run("Culinary 2")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Week 1 School Working Copy")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    date = doc.add_paragraph("August 11-14, 2026 | Four instructional days")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Included", level=2)
    for item in [
        "Four Daily Teaching Guides",
        "Active projection source",
        "Required group recipe and mise plan",
        "Instructor/TA observation and reusable support tools",
        "End-of-day curriculum capture prompts",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    note = doc.add_paragraph()
    add_rich_text(note, "**School workflow:** Upload this DOCX to the work Google Drive and open it as a Google Doc when edits are needed. Print the companion group-recipe PDF; print other supports only when useful.")
    shade_paragraph(note, "DDEBF7")
    doc.add_page_break()
    for filename in TEACHER:
        add_markdown(doc, SRC / filename)
        doc.add_page_break()
    add_divider(doc, "Instructional Supports")
    for filename in SUPPORTS:
        add_markdown(doc, SRC / "Student_Materials" / filename)
        doc.add_page_break()
    add_markdown(doc, SRC / "Slides" / "Culinary_2_Week_01_Slide_Source.md")
    doc.core_properties.title = "Culinary 2 Week 1 School Working Copy"
    doc.core_properties.author = "WACOS"
    doc.save(path)


def build_recipe_pdf(path):
    sample = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=sample["Title"], fontName="Helvetica-Bold", fontSize=18, leading=21, textColor=colors.HexColor("#17365D"), spaceAfter=4)
    sub = ParagraphStyle("sub", parent=sample["BodyText"], fontName="Helvetica-Bold", fontSize=8.5, leading=10.5, spaceAfter=4)
    heading = ParagraphStyle("heading", parent=sample["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=13, textColor=colors.HexColor("#2F75B5"), spaceBefore=4, spaceAfter=2)
    body = ParagraphStyle("body", parent=sample["BodyText"], fontName="Helvetica", fontSize=8.4, leading=10.3, spaceAfter=1.5)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=13, firstLineIndent=-8)
    document = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.62 * inch, rightMargin=0.62 * inch, topMargin=0.48 * inch, bottomMargin=0.48 * inch, title="Culinary 2 Week 1 Vegetable Rice Bowl Recipe")
    story = [
        Paragraph("Vegetable Rice Bowl - Week 1 Group Recipe", title),
        Paragraph("Culinary 2 | Tuesday planning and Friday baseline lab", sub),
        Paragraph("Planned yield: 3 portions &nbsp;&nbsp;&nbsp; Actual yield: __________ portions", sub),
        Paragraph("Ingredients", heading),
    ]
    ingredients = [
        "1 cup long-grain rice",
        "2 cups water",
        "1/2 teaspoon salt for the rice, or instructor-approved amount",
        "1 tablespoon vegetable oil, divided as needed",
        "1/2 onion; 1 carrot; 1 celery stalk",
        "1/2 bell pepper",
        "Salt and pepper; optional instructor-approved acid or fresh herb",
    ]
    story.extend(Paragraph("- " + item, bullet) for item in ingredients)
    story.append(Paragraph("Method", heading))
    method = [
        "Read the complete recipe, confirm group roles, and post the flames-off time.",
        "Wash hands and prepare the approved sanitation system.",
        "Gather equipment and ingredients. Complete mise en place before heat.",
        "Rinse or do not rinse the rice according to the instructor's direction.",
        "Combine rice, 2 cups water, and 1/2 teaspoon salt in a lidded stovetop pot.",
        "Bring to a boil. Stir once, cover, reduce heat to low, and cook 15 minutes. Do not lift the lid.",
        "Remove from heat and rest, covered, for 10 minutes. Fluff with a fork.",
        "Retrieve Wednesday's labeled vegetables after instructor approval; complete any additional cutting only with authorization.",
        "Heat a saute pan over medium to medium-high heat. Add oil, then onion, carrot, celery, and bell pepper.",
        "Saute approximately 5-8 minutes until tender with retained color and texture. Communicate before quality is at risk.",
        "Taste safely and make one purposeful seasoning adjustment.",
        "Portion three consistent bowls; record actual yield and any approved change.",
        "Store or discard food as directed and complete the full station reset.",
    ]
    story.extend(Paragraph(f"{number}. {item}", body) for number, item in enumerate(method, 1))
    story.extend([
        Paragraph("Mark the Recipe on Tuesday", heading),
        Paragraph("Circle yield, time, heat directions, and measurements. Box equipment. Star action verbs. Underline hazards or approval points. Arrow quality cues. Mark the first three actions and shutdown plan.", body),
        PageBreak(),
        Paragraph("Group Mise and Timing Plan", title),
        Paragraph("Lead cook: ____________________ &nbsp;&nbsp; Mise cook: ____________________ &nbsp;&nbsp; Sanitation/service cook: ____________________", body),
        Spacer(1, 8),
        Paragraph("First three actions", heading),
        Paragraph("1. __________________________________________________________________________", body),
        Paragraph("2. __________________________________________________________________________", body),
        Paragraph("3. __________________________________________________________________________", body),
        Spacer(1, 8),
        Paragraph("Equipment", heading),
        Paragraph("____________________________________________________________________________", body),
        Paragraph("____________________________________________________________________________", body),
        Paragraph("Hazards or approval points", heading),
        Paragraph("____________________________________________________________________________", body),
        Paragraph("____________________________________________________________________________", body),
        Paragraph("Quality cues", heading),
        Paragraph("____________________________________________________________________________", body),
        Paragraph("____________________________________________________________________________", body),
        Spacer(1, 10),
        Paragraph("Flames off: __________ &nbsp;&nbsp;&nbsp; Reset complete: __________", sub),
        Spacer(1, 16),
        Paragraph("Keep this marked copy with the group for Friday. Heat and equipment stay off until instructor approval.", sub),
    ])
    document.build(story)


def build():
    PKG.mkdir(parents=True, exist_ok=True)
    for item in PKG.iterdir():
        if item.is_file():
            item.unlink()
    build_docx(PKG / "01_Culinary_2_Week_01_Editable.docx")
    build_recipe_pdf(PKG / "02_Culinary_2_Week_01_Required_Print_Recipe_and_Mise_Plan.pdf")
    (PKG / "README.txt").write_text(
        "CULINARY 2 - WEEK 1 SCHOOL UPLOAD PACKAGE\n\n"
        "01 Editable DOCX: four Daily Teaching Guides, active projection source, TA observation tools, reusable group supports, and curriculum-capture prompts.\n"
        "02 Required Print PDF: print one Vegetable Rice Bowl Recipe and Mise Plan per group; use Tuesday and Friday.\n\n"
        "The vegetable rice bowl is the approved Week 1 vehicle. Rice uses the stovetop method; usable Wednesday vegetables are safely labeled, refrigerated, and carried forward for Friday.\n"
        "Other support pages may be projected, reused, or printed selectively.\n",
        encoding="utf-8",
    )
    if ZIP.exists():
        ZIP.unlink()
    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as archive:
        for item in sorted(PKG.iterdir()):
            archive.write(item, arcname=f"Culinary_2_Week_01/{item.name}")
    print(PKG)
    print(ZIP)


if __name__ == "__main__":
    build()
